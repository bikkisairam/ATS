from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# -----------------------------
# Style Helpers
# -----------------------------
def _set_style(paragraph, size=10, bold=False):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.bold = bold


def _add_heading(doc, text):
    p = doc.add_paragraph(text.upper())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_style(p, size=11, bold=True)


def _add_text(doc, text):
    p = doc.add_paragraph(text)
    _set_style(p, size=10)


def _add_bullets(doc, bullets):
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        _set_style(p, size=10)


# -----------------------------
# DOCX Builder
# -----------------------------
def build_docx(resume: dict, output_path: Path):
    doc = Document()

    # -----------------------------
    # Header
    # -----------------------------
    details = resume.get("Details", {})
    name = details.get("Name", "")
    contact = " | ".join(
        v for v in [
            details.get("Email"),
            details.get("Phone"),
            details.get("Location"),
            details.get("LinkedIn"),
            details.get("GitHub"),
        ] if v
    )

    p_name = doc.add_paragraph(name)
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_style(p_name, size=16, bold=True)

    if contact:
        p_contact = doc.add_paragraph(contact)
        p_contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _set_style(p_contact, size=9)

    # -----------------------------
    # Summary
    # -----------------------------
    summary = resume.get("Summary")
    if summary:
        _add_heading(doc, "Summary")
        _add_text(doc, summary)

    # -----------------------------
    # Skills
    # -----------------------------
    skills = resume.get("Skills", [])
    if skills:
        _add_heading(doc, "Skills")
        _add_text(doc, ", ".join(skills))

    # -----------------------------
    # Experience
    # -----------------------------
    experience = resume.get("Work Experience", [])
    if experience:
        _add_heading(doc, "Experience")
        for job in experience:
            title = f"{job.get('Company Name', '')} — {job.get('Role', '')}"
            date = job.get("Date", "")
            _add_text(doc, f"{title} ({date})" if date else title)
            _add_bullets(doc, job.get("Bullet Points", []))

    # -----------------------------
    # Projects
    # -----------------------------
    projects = resume.get("Project Experience", [])
    if projects:
        _add_heading(doc, "Projects")
        for proj in projects:
            header = proj.get("Title", "")
            tech = proj.get("Tech Stack")
            _add_text(doc, f"{header} ({tech})" if tech else header)
            _add_bullets(doc, proj.get("Bullet Points", []))

    # -----------------------------
    # Education
    # -----------------------------
    education = resume.get("Education", [])
    if education:
        _add_heading(doc, "Education")
        for edu in education:
            line = f"{edu.get('Degree', '')}, {edu.get('Institution', '')}"
            date = edu.get("Date", "")
            _add_text(doc, f"{line} ({date})" if date else line)

    # -----------------------------
    # Certifications
    # -----------------------------
    certs = resume.get("Achievements and Certifications", [])
    if certs:
        _add_heading(doc, "Achievements & Certifications")
        _add_bullets(doc, certs)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
